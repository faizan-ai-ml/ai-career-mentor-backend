import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_polished_report():
    doc = Document()

    # --- Configuration: Image Paths ---
    BASE_DIR = r"C:\Users\DELL\.gemini\antigravity\brain\eaa3fac7-e23b-412d-b87e-84b8df8eabb3"
    IMG_ARCH = os.path.join(BASE_DIR, "system_architecture_diagram_1767641640518.png")
    IMG_USE_CASE = os.path.join(BASE_DIR, "use_case_diagram_fyp_1767641655815.png")
    IMG_ERD = os.path.join(BASE_DIR, "database_erd_visualization_1767641671595.png")
    IMG_UI = os.path.join(BASE_DIR, "mobile_ui_mockup_roadmap_1767641689367.png")

    # --- Title Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI CAREER MENTOR\n')
    run.bold = True
    run.font.size = Pt(36)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A Cross-Disciplinary Automated Career Guidance System\n')
    run.font.size = Pt(18)
    run.italic = True

    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Final Year Project Report\n')
    run.font.size = Pt(22)
    run.bold = True

    doc.add_paragraph('\n' * 3)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted to:\nDepartment of Computer Science & Software Engineering\n\nSubmitted by:\n[Your Name]\nRegistration No: [Your ID]')
    run.font.size = Pt(14)
    
    doc.add_paragraph('\n' * 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Date: 2026')

    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "Chapter 1: Introduction",
        "  1.1 Background",
        "  1.2 Problem Statement",
        "  1.3 Objectives",
        "  1.4 Scope",
        "Chapter 2: Existing System & Proposed System",
        "  2.1 Existing System",
        "  2.2 Limitations of Existing System",
        "  2.3 Proposed System",
        "  2.4 Advantages of Proposed System",
        "Chapter 3: System Analysis",
        "  3.1 Functional Requirements",
        "  3.2 Non-Functional Requirements",
        "  3.3 Use Case Diagram",
        "  3.4 Use Case Descriptions",
        "Chapter 4: System Design",
        "  4.1 System Architecture",
        "  4.2 Database Design (ERD)",
        "  4.3 API Design",
        "  4.4 User Interface Design",
        "Chapter 5: Implementation",
        "  5.1 Tools & Technologies",
        "  5.2 Backend Implementation",
        "  5.3 Mobile App (Android XML) Implementation",
        "  5.4 AI Integration Module",
        "Chapter 6: Testing",
        "  6.1 Testing Strategy",
        "  6.2 Test Case Tables",
        "  6.3 Result Summary",
        "Chapter 7: Conclusion & Future Enhancements",
        "  References"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        if "Chapter" in item:
            p.runs[0].bold = True
    
    doc.add_page_break()

    # --- Chapter 1 ---
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph(
        "Career navigation is a complex, lifelong journey that affects individuals across all professional domains. "
        "Whether a student is entering the workforce or a seasoned professional is pivoting to a new industry, "
        "the need for personalized guidance is universal. Traditional career mentorship relies on human intervention, "
        "which is often expensive, subjective, and limited by geographical availability. "
        "With the advent of Large Language Models (LLMs), there is a unique opportunity to provide automated, "
        "high-quality career counseling to the general public regardless of their field—be it Tech, Business, Healthcare, or Arts."
    )
    
    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph(
        "Modern job seekers face a 'transparency gap'. Resumes are often static and fail to communicate potential, "
        "while recruiters in different sectors have highly specific, shifting expectations. Individuals lack a "
        "mechanism to objectively measure their skills against global standards. Furthermore, finding a structured "
        "path to bridge these gaps is a daunting task, often leading to 'information overload' where users "
        "feel paralyzed by the sheer volume of available online courses and certifications."
    )

    doc.add_heading('1.3 Objectives', level=2)
    objs = [
        "To develop a universal career mentorship mobile app using Kotlin and Android XML.",
        "To create a scalable FastAPI backend for cross-disciplinary resume analysis.",
        "To integrate state-of-the-art AI (DeepSeek/Gemini) for semantic skill extraction.",
        "To automate the generation of personalized 8-week learning roadmaps.",
        "To provide a real-time AI consultant for 24/7 career support."
    ]
    for obj in objs:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('1.4 Scope', level=2)
    doc.add_paragraph(
        "The project covers the end-to-end development of the AI Career Mentor platform. This includes "
        "user authentication, secure file processing, AI-driven gap analysis, and result persistence. "
        "The scope is designed to be inclusive, supporting professionals from both technical and non-technical "
        "backgrounds, and providing guidance that is tailored to their specific career goals."
    )

    doc.add_page_break()

    # --- Chapter 2 ---
    doc.add_heading('Chapter 2: Existing System & Proposed System', level=1)
    doc.add_heading('2.1 Existing System', level=2)
    doc.add_paragraph(
        "The current ecosystem is dominated by manual career centers and static job boards. "
        "University-led centers provide human interaction but suffer from high 'student-to-mentor' ratios. "
        "Platforms like LinkedIn and Indeed primarily focus on matching existing skills to jobs without "
        "offering a developmental roadmap for improvement."
    )
    
    doc.add_heading('2.2 Limitations', level=2)
    doc.add_paragraph("1. Lack of scalability for global audiences.")
    doc.add_paragraph("2. High cost of professional mentorship services.")
    doc.add_paragraph("3. Difficulty in identifying 'hidden' skill gaps that aren't apparent to the user.")
    
    doc.add_heading('2.3 Proposed System', level=2)
    doc.add_paragraph(
        "The proposed AI Career Mentor offers an automated, semantic approach. By analyzing the 'meaning' "
        "behind a user's experience rather than just keywords, the system provides a far more accurate "
        "representation of their professional standing. It suggests a curated path to success, effectively "
        "acting as a virtual project manager for the user's career development."
    )
    
    doc.add_heading('2.4 Advantages', level=2)
    doc.add_paragraph("Scalable, Objective, Always-on, and Hyper-personalized learning path generation.")

    doc.add_page_break()

    # --- Chapter 3 ---
    doc.add_heading('Chapter 3: System Analysis', level=1)
    doc.add_heading('3.1 Functional Requirements', level=2)
    doc.add_paragraph("F1: The system must allow users to register and login securely.")
    doc.add_paragraph("F2: The system must support PDF resume uploads for processing.")
    doc.add_paragraph("F3: The AI module must extract skills and compare them to the target role.")
    doc.add_paragraph("F4: The app must display a week-by-week learning roadmap with resources.")
    doc.add_paragraph("F5: A real-time chat interface must handle career-related user queries.")
    
    doc.add_heading('3.2 Non-Functional Requirements', level=2)
    doc.add_paragraph("N1: Performance - Analysis completed in under 25 seconds.")
    doc.add_paragraph("N2: Security - Passwords stored as Bcrypt hashes; tokens refreshed via JWT.")
    doc.add_paragraph("N3: Scalability - Support for 50+ concurrent analysis requests.")

    doc.add_heading('3.3 Use Case Diagram', level=2)
    if os.path.exists(IMG_USE_CASE):
        doc.add_picture(IMG_USE_CASE, width=Inches(4.5))
        p = doc.add_paragraph("Figure 3.1: System Use Case Diagram")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Use Case Diagram Placeholder]")

    doc.add_heading('3.4 Use Case Description', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Use Case'; hdr[1].text = 'Actor'; hdr[2].text = 'Main Success Scenario'
    
    cases = [
        ('Analyze Career', 'User', 'User uploads PDF, AI generates roadmap and saves to profile'),
        ('Consult AI', 'User', 'User sends query in chat, AI returns tailored career advice'),
        ('View Timeline', 'User', 'User accesses dashboard to see progress on their 8-week plan')
    ]
    for n, a, d in cases:
        row = table.add_row().cells
        row[0].text=n; row[1].text=a; row[2].text=d

    doc.add_page_break()

    # --- Chapter 4 ---
    doc.add_heading('Chapter 4: System Design', level=1)
    doc.add_heading('4.1 System Architecture', level=2)
    if os.path.exists(IMG_ARCH):
        doc.add_picture(IMG_ARCH, width=Inches(4.5))
        p = doc.add_paragraph("Figure 4.1: High-Level System Architecture")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "The system architecture is divided into three layers: Presentation (Android), Business Logic (FastAPI), "
        "and Knowledge Layer (AI Services). Communication is managed via JSON-based REST APIs."
    )
    
    doc.add_heading('4.2 Database Design (ERD)', level=2)
    if os.path.exists(IMG_ERD):
        doc.add_picture(IMG_ERD, width=Inches(4.0))
        p = doc.add_paragraph("Figure 4.2: Entity Relationship Diagram")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("The database follows a normalized relational structure using PostgreSQL.")

    doc.add_heading('4.3 API Design', level=2)
    doc.add_paragraph("Key API Routes:\n- POST /upload: Handles file streaming and triggers AI parser.\n- GET /roadmap: Fetches the personalized JSON object.")

    doc.add_heading('4.4 UI Design', level=2)
    if os.path.exists(IMG_UI):
        doc.add_picture(IMG_UI, width=Inches(3.0))
        p = doc.add_paragraph("Figure 4.3: App UI Design Mockup")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Uses Material 3 Dark theme with clean, minimalist navigation.")

    doc.add_page_break()

    # --- Chapter 5 ---
    doc.add_heading('Chapter 5: Implementation', level=1)
    doc.add_heading('5.1 Tools & Technologies', level=2)
    doc.add_paragraph("Backend: Python (FastAPI), SQLAlchemy, Pydantic, Gunicorn.\nFrontend: Kotlin, XML, Retrofit, View Binding.\nAI: DeepSeek-V3, Google Gemini-1.5-Pro.")
    
    doc.add_heading('5.2 Backend Implementation', level=2)
    doc.add_paragraph(
        "The server is implemented as an asynchronous FastAPI application. It uses a custom 'Service Layer' "
        "to interact with AI providers. Large PDF files are processed using 'pypdf' which converts them to "
        "structured text strings before analysis."
    )
    
    doc.add_heading('5.3 Mobile App Implementation', level=2)
    doc.add_paragraph(
        "The Android application strictly follows the MVVM pattern. User Interface is defined in standard "
        "XML layout files to ensure maximum performance. Networking is handled via Retrofit with "
        "coroutines for non-blocking I/O operations."
    )

    doc.add_heading('5.4 AI Integration Module', level=2)
    doc.add_paragraph(
        "AI prompts are engineered to return 'Structured JSON' using system-level instructions. This ensures "
        "that the mobile app can dynamically render the learning roadmap without parsing raw text."
    )

    doc.add_page_break()

    # --- Chapter 6 ---
    doc.add_heading('Chapter 6: Testing', level=1)
    doc.add_heading('6.1 Testing Strategy', level=2)
    doc.add_paragraph("Black box testing was primarily used for functional verification, while postman was used for API endpoint testing.")
    
    doc.add_heading('6.2 Test Case Tables', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Test ID'; hdr[1].text = 'Case'; hdr[2].text = 'Input'; hdr[3].text = 'Expectation'; hdr[4].text = 'Result'
    
    tc_data = [
        ('TC-01', 'Auth Success', 'Correct User/Pass', 'Access Granted', 'Pass'),
        ('TC-02', 'File Type', 'Image File', 'Error Message', 'Pass'),
        ('TC-03', 'Roadmap', 'Non-tech Resume', 'Art-based Roadmp', 'Pass'),
        ('TC-04', 'Chat Latency', 'Fast Query', 'Response < 5s', 'Pass')
    ]
    for tid, c, i, e, r in tc_data:
        row = table.add_row().cells
        row[0].text=tid; row[1].text=c; row[2].text=i; row[3].text=e; row[4].text=r

    doc.add_page_break()

    # --- Chapter 7 ---
    doc.add_heading('Chapter 7: Conclusion & Future Enhancements', level=1)
    doc.add_paragraph(
        "The AI Career Mentor represents a significant step toward making professional guidance "
        "equitable and available to all. By leveraging cross-disciplinary AI, the project bridges "
        "the gap between human ambition and industry reality."
    )
    doc.add_heading('Future Enhancements', level=2)
    doc.add_paragraph("- Integration with LinkedIn API for job discovery.\n- Real-time voice mentorship features.\n- Interactive coding/design challenges within the app.")

    doc.add_heading('References', level=2)
    doc.add_paragraph("1. Python FastAPI Documentation - https://fastapi.tiangolo.com/\n2. Android Material Design Guidelines - https://m3.material.io/\n3. OpenRouter AI Models - https://openrouter.ai/")

    doc_path = 'AI_Career_Mentor_Final_Polished_Report.docx'
    doc.save(doc_path)
    print(f"Polished Report generated: {doc_path}")

if __name__ == "__main__":
    create_final_polished_report()
