import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_fyp_word_report():
    doc = Document()

    # --- Title Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI CAREER MENTOR')
    run.bold = True
    run.font.size = Pt(28)
    
    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Final Year Project Report')
    run.font.size = Pt(20)

    doc.add_paragraph('\n' * 3)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Submitted to: [Department Name]\n[University Name]\n\nSubmitted by: [Your Name]\nRegistration No: [Your ID]')
    
    doc.add_paragraph('\n' * 5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Date: 2026')

    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "Chapter 1: Introduction",
        "  1.1 Background",
        "  1.2 Problem Statement",
        "  1.3 Objectives",
        "  1.4 Scope",
        "Chapter 2: Existing System & Proposed System",
        "  2.1 Existing System",
        "  2.2 Limitations of Existing System",
        "  2.3 Proposed System",
        "  2.4 Advantages of Proposed System",
        "Chapter 3: System Analysis",
        "  3.1 Functional Requirements",
        "  3.2 Non-Functional Requirements",
        "  3.3 Use Case Diagram",
        "  3.4 Use Case Descriptions",
        "Chapter 4: System Design",
        "  4.1 System Architecture",
        "  4.2 Database Design (ERD)",
        "  4.3 API Design",
        "  4.4 User Interface Design",
        "Chapter 5: Implementation",
        "  5.1 Tools & Technologies",
        "  5.2 Backend Implementation",
        "  5.3 Mobile App Implementation",
        "  5.4 AI Career Recommendation Module",
        "Chapter 6: Testing",
        "  6.1 Testing Strategy",
        "  6.2 Test Cases",
        "  6.3 Result Summary",
        "Chapter 7: Conclusion & Future Enhancements",
        "  References"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()

    # --- Chapter 1 ---
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph(
        "Career transitions and industry entries are critical phases for individuals regardless of their domain. "
        "Artificial Intelligence (AI), particularly LLMs, can now perform semantic analysis across any discipline. "
        "This project explores universal career mentorship by developing a system that analyzes both technical and non-technical resumes to generate actionable roadmaps for any career path."
    )
    
    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph(
        "Professionals across all sectors lack objective resume feedback and face difficulty identifying skill gaps when aiming for pivots or promotions. "
        "The lack of a structured, cross-disciplinary learning path creates a significant barrier to career growth."
    )

    doc.add_heading('1.3 Objectives', level=2)
    objs = [
        "To develop a universal native Android application for career mentorship accessible to all.",
        "To implement a robust FastAPI backend for multi-disciplinary processing.",
        "To automate resume analysis for both technical and non-technical domains using LLMs.",
        "To generate personalized weekly roadmaps for any industry niche."
    ]
    for obj in objs:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('1.4 Scope', level=2)
    doc.add_paragraph("Includes Mobile App, Backend API, AI Orchestration, and domain-agnostic roadmap generation for the general public.")

    doc.add_page_break()

    # --- Chapter 2 ---
    doc.add_heading('Chapter 2: Existing System & Proposed System', level=1)
    doc.add_heading('2.1 Existing System', level=2)
    doc.add_paragraph("Relies on manual university centers and static portals like LinkedIn/Indeed.")
    
    doc.add_heading('2.2 Limitations', level=2)
    doc.add_paragraph("Manual counseling lacks scalability, availability, and hyper-personalization.")
    
    doc.add_heading('2.3 Proposed System', level=2)
    doc.add_paragraph("An AI-driven platform providing 24/7 semantic gap analysis and automated career paths.")
    
    doc.add_heading('2.4 Advantages', level=2)
    doc.add_paragraph("Precision, 24/7 availability, and unique customized outputs for every student.")

    doc.add_page_break()

    # --- Chapter 3 ---
    doc.add_heading('Chapter 3: System Analysis', level=1)
    doc.add_heading('3.1 Functional Requirements', level=2)
    doc.add_paragraph("Includes Auth, PDF Upload, Skill Extraction, Roadmap View, and AI Chatbot.")
    
    doc.add_heading('3.2 Non-Functional Requirements', level=2)
    doc.add_paragraph("Covers Performance (<30s analysis), Security (Bcrypt/JWT), and Availability.")

    doc.add_heading('3.3 Use Case Diagram', level=2)
    doc.add_paragraph("[PLEASE INSERT USE CASE DIAGRAM IMAGE HERE]")

    doc.add_heading('3.4 Use Case Description', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Use Case'
    hdr_cells[1].text = 'Actor'
    hdr_cells[2].text = 'Description'
    
    cases = [
        ('Login', 'Student', 'Authenticates with email/password'),
        ('Upload Resume', 'Student', 'Upload PDF for AI analysis'),
        ('Chat with AI', 'Student', 'Ask career questions to LLM')
    ]
    for name, actor, desc in cases:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = actor
        row[2].text = desc

    doc.add_page_break()

    # --- Chapter 4 ---
    doc.add_heading('Chapter 4: System Design', level=1)
    doc.add_heading('4.1 System Architecture', level=2)
    doc.add_paragraph("Distributed Client-Server architecture using Kotlin (Mobile) and FastAPI (Server).")
    
    doc.add_heading('4.2 Database Design (ERD)', level=2)
    doc.add_paragraph("Highly normalized schema with Users and StudentProfiles tables.")
    doc.add_paragraph("[PLEASE INSERT ER DIAGRAM IMAGE HERE]")

    doc.add_heading('4.3 API Design', level=2)
    doc.add_paragraph("REST API endpoints for Authentication, Profile, and Resume Analysis.")

    doc.add_heading('4.4 UI Design', level=2)
    doc.add_paragraph("Designed using standard Android XML layouts for high compatibility.")

    doc.add_page_break()

    # --- Chapter 5 ---
    doc.add_heading('Chapter 5: Implementation', level=1)
    doc.add_heading('5.1 Tools & Technologies', level=2)
    doc.add_paragraph("Kotlin, Python, FastAPI, PostgreSQL, Retrofit, Hilt, DeepSeek/Gemini.")
    
    doc.add_heading('5.2 Backend', level=2)
    doc.add_paragraph("Async Python processing with SQLAlchemy for Postgres database management.")
    
    doc.add_heading('5.3 Mobile App', level=2)
    doc.add_paragraph("MVVM pattern with ViewBinding and Retrofit for network calls.")

    doc.add_heading('5.4 AI Integration', level=2)
    doc.add_paragraph("Orchestration module to handle LLM prompt construction and JSON parsing.")

    doc.add_page_break()

    # --- Chapter 6 ---
    doc.add_heading('Chapter 6: Testing', level=1)
    doc.add_heading('6.1 Strategy', level=2)
    doc.add_paragraph("Module testing followed by system integration and physical device testing.")
    
    doc.add_heading('6.2 Test Cases', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'ID'; hdr[1].text = 'Scenario'; hdr[2].text = 'Input'; hdr[3].text = 'Expected'; hdr[4].text = 'Status'
    
    tests = [
        ('TC-01', 'Signup', 'New User Data', 'Account Created', 'Passed'),
        ('TC-02', 'Analysis', 'PDF Resume', 'Roadmap JSON', 'Passed'),
        ('TC-03', 'Chat', 'Query', 'AI Advice', 'Passed')
    ]
    for tid, scn, inp, exp, stat in tests:
        row = table.add_row().cells
        row[0].text=tid; row[1].text=scn; row[2].text=inp; row[3].text=exp; row[4].text=stat

    doc.add_page_break()

    # --- Chapter 7 ---
    doc.add_heading('Chapter 7: Conclusion', level=1)
    doc.add_paragraph("The AI Career Mentor empowers students to bridge the skill gap using automated mentorship.")
    doc.add_heading('References', level=2)
    doc.add_paragraph("FastAPI Documentation, Android Developers Guide, OpenRouter/Gemini API Docs.")

    doc_path = 'AI_Career_Mentor_Final_Report_V2.docx'
    doc.save(doc_path)
    print(f"Report generated: {doc_path}")

if __name__ == "__main__":
    create_fyp_word_report()
