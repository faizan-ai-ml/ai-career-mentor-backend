import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    document = Document()

    # Title Page
    document.add_heading('AI CAREER MENTOR', 0)
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FINAL YEAR PROJECT REPORT')
    run.bold = True
    run.font.size = Pt(16)
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\nSubmitted by: [Your Name]\nSubmission Date: 2026\nDomain: Artificial Intelligence & Mobile App Development')

    document.add_page_break()

    # Table of Contents (Placeholder text as real TOC is complex in python-docx)
    document.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Abstract",
        "2. Chapter 1: Introduction",
        "3. Chapter 2: Proposed Solution",
        "4. Chapter 3: System Analysis",
        "5. Chapter 4: System Design",
        "6. Chapter 5: Implementation Details",
        "7. Chapter 6: User Interface",
        "8. Chapter 7: Conclusion & Future Scope"
    ]
    for item in toc_items:
        document.add_paragraph(item)
    
    document.add_page_break()

    # 1. Abstract
    document.add_heading('1. Abstract', level=1)
    document.add_paragraph(
        "The AI Career Mentor is an innovative mobile application designed to democratize career counseling for university students. "
        "By utilizing state-of-the-art Generative AI (Large Language Models), the system automates the process of resume review and career planning. "
        "Unlike traditional static platforms, this system offers hyper-personalized, dynamic learning roadmaps based on a deep semantic analysis of the student's current skills versus industry demands. "
        "This report details the design, development, and implementation of the system using Python FastAPI and native Android technologies."
    )

    # 2. Introduction
    document.add_heading('2. Chapter 1: Introduction', level=1)
    
    document.add_heading('1.1 Problem Background', level=2)
    document.add_paragraph(
        "In the current educational landscape, there is a significant disconnect between academic curriculum and industry expectations. "
        "Students often graduate with theoretical knowledge but lack specific, practical skills required by employers. "
        "Professional career counseling is often expensive or inaccessible."
    )

    document.add_heading('1.2 Problem Statement', level=2)
    p = document.add_paragraph()
    run = p.add_run('"Students lack accurate, real-time feedback on their resumes and struggle to create structured, actionable learning paths to bridge their skill gaps."')
    run.italic = True
    
    document.add_heading('1.3 Objectives', level=2)
    objectives = [
        "To develop an automated system for extracting and analyzing detailed resume data from PDF files.",
        "To utilize AI to identify skill deficiencies compared to current market standards.",
        "To generate customized, week-by-week learning roadmaps including curated resources.",
        "To provide an accessible mobile interface for tracking progress."
    ]
    for obj in objectives:
        document.add_paragraph(obj, style='List Bullet')

    document.add_heading('1.4 Project Scope', level=2)
    document.add_paragraph(
        "The project covers the development of a RESTful API backend, a native Android client, and the integration of third-party AI services."
    )

    # 3. Proposed Solution
    document.add_heading('3. Chapter 2: Proposed Solution', level=1)
    
    document.add_heading('2.1 Solution Overview', level=2)
    document.add_paragraph(
        "We propose an AI-powered platform where a student simply uploads their resume. "
        "The system 'reads' the resume, understands the student's career goal (e.g., 'Python Developer'), and acts as a virtual mentor."
    )

    document.add_heading('2.2 Key Features', level=2)
    features = [
        "Smart Resume Parser: Converts raw PDF data into structured JSON profiles.",
        "Gap Analysis Engine: Compares User Skills vs. Target Role requirements.",
        "Dynamic Roadmap Generator: Creates a unique study plan tailored to the user's missing skills.",
        "24/7 AI Chatbot: A context-aware chat interface for career advice."
    ]
    for f in features:
        document.add_paragraph(f, style='List Bullet')

    # 4. System Analysis
    document.add_heading('4. Chapter 3: System Analysis', level=1)
    
    document.add_heading('3.1 Functional Requirements', level=2)
    document.add_paragraph("User Module: Signup, Login (JWT), Profile Management.", style='List Bullet')
    document.add_paragraph("Analysis Module: File Upload, Resume Parsing, AI Processing.", style='List Bullet')
    document.add_paragraph("Roadmap Module: View Roadmap, Mark Topics as Complete.", style='List Bullet')

    document.add_heading('3.2 Non-Functional Requirements', level=2)
    document.add_paragraph("Performance: Analysis < 30 seconds.", style='List Bullet')
    document.add_paragraph("Security: Encrypted data, hashed passwords.", style='List Bullet')
    
    document.add_heading('3.3 Technology Stack', level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Reason'
    
    tech_data = [
        ('Frontend', 'Android Native (Kotlin)', 'Performance'),
        ('UI Design', 'XML Layouts', 'Standard, compatible UI'),
        ('Backend', 'Python FastAPI', 'High performance, Async'),
        ('Database', 'PostgreSQL', 'Relational integrity'),
        ('AI Model', 'DeepSeek V3 / Gemini', 'Reasoning capability')
    ]
    for component, tech, reason in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech
        row_cells[2].text = reason

    # 5. System Design
    document.add_heading('5. Chapter 4: System Design', level=1)
    
    document.add_heading('4.1 System Architecture', level=2)
    document.add_paragraph(
        "The system follows a Client-Server Architecture. The Android Client uses MVVM (Model-View-ViewModel) to separate logic from XML views. "
        "The Backend exposes REST endpoints and orchestrates AI services."
    )
    
    document.add_heading('4.2 Database Design', level=2)
    document.add_paragraph("Key tables include:")
    document.add_paragraph("Users: Stores secure credentials.", style='List Bullet')
    document.add_paragraph("StudentProfiles: Stores skills and roadmap JSON data.", style='List Bullet')

    # 6. Implementation
    document.add_heading('6. Chapter 5: Implementation Details', level=1)
    
    document.add_heading('5.1 Backend Implementation', level=2)
    document.add_paragraph("Built with FastAPI using Dependency Injection. Auth implemented via OAuth2PasswordBearer and PyJWT.")
    
    document.add_heading('5.2 Android Implementation', level=2)
    document.add_paragraph("Written in Kotlin. Uses XML Layouts for UI, Retrofit for networking, and Hilt for dependency injection.")

    # 7. User Interface
    document.add_heading('7. Chapter 6: User Interface', level=1)
    document.add_paragraph("[Paste Screenshots Here for Login, Dashboard, Analysis, and Roadmap]")

    # 8. Conclusion
    document.add_heading('8. Chapter 7: Conclusion & Future Scope', level=1)
    document.add_paragraph(
        "The AI Career Mentor Project demonstrates the practical application of AI in education. "
        "It provides a scalable, automated solution for career guidance."
    )
    document.add_paragraph("Future Scope: Mock Interviews, Job Matching integration.")

    document.save('Project_Documentation.docx')
    print("Document created successfully at Project_Documentation.docx")

if __name__ == "__main__":
    create_document()
